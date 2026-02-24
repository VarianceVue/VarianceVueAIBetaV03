"""
Stage 1 extractors: PDF, DOCX, XLSX, TXT/CSV, XER (Primavera P6 export).
Reuses existing PDF extraction; adds DOCX, XLSX, and XER (text decode).
"""
from __future__ import annotations

import io
from typing import Tuple


def _extract_pdf(bytes_in: bytes) -> str:
    from schedule_agent_web.upload_utils import extract_text_from_pdf
    text = extract_text_from_pdf(bytes_in)
    text = (text or "").strip()
    if not text:
        return (
            "[This PDF could not be read as text. It may be image-only (scanned). "
            "Use an OCR tool or upload a text-based PDF.]"
        )
    return text


def _extract_docx(bytes_in: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(bytes_in))
        parts = []
        for para in doc.paragraphs:
            t = (para.text or "").strip()
            if t:
                parts.append(t)
        for table in doc.tables:
            for row in table.rows:
                cells = [str(cell.text or "").strip() for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n\n".join(parts) if parts else ""
    except Exception:
        return ""


def _extract_xlsx(bytes_in: bytes) -> str:
    try:
        from openpyxl import load_workbook
        wb = load_workbook(io.BytesIO(bytes_in), read_only=True, data_only=True)
        parts = []
        for sheet in wb.worksheets:
            parts.append(f"--- Sheet: {sheet.title} ---")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c or "").strip() for c in (row or [])]
                if any(cells):
                    parts.append("\t".join(cells))
        wb.close()
        return "\n".join(parts) if parts else ""
    except Exception:
        return ""


def _extract_txt(bytes_in: bytes) -> str:
    try:
        return bytes_in.decode("utf-8", errors="replace").strip()
    except Exception:
        try:
            return bytes_in.decode("latin-1", errors="replace").strip()
        except Exception:
            return ""


def _extract_xer(bytes_in: bytes) -> str:
    """XER is Primavera P6 export — text-based; decode as UTF-8 or Latin-1."""
    return _extract_txt(bytes_in)


def _extract_image(bytes_in: bytes, filename: str) -> str:
    """
    Use Claude vision to describe the image (e.g. site photo). Requires ANTHROPIC_API_KEY.
    Returns description text for storage/vectorization; fallback message if vision unavailable.
    """
    ext = (filename or "").lower().split(".")[-1] if "." in (filename or "") else ""
    media_map = {"jpg": "image/jpeg", "jpeg": "image/jpeg", "png": "image/png", "gif": "image/gif", "webp": "image/webp"}
    media_type = media_map.get(ext, "image/jpeg")
    try:
        from schedule_agent_web.vision import describe_image
        text, err = describe_image(bytes_in, media_type=media_type)
        if err:
            return f"[Image: vision unavailable — {err}. Upload with ANTHROPIC_API_KEY set for automatic description.]"
        return text or "[Image: no description returned.]"
    except Exception as e:
        return f"[Image: vision error — {e}. Set ANTHROPIC_API_KEY for site photo descriptions.]"


def detect_format(filename: str, bytes_in: bytes) -> str:
    """Return normalized format: pdf, docx, xlsx, txt, csv, xer, image, or unknown."""
    ext = (filename or "").lower().split(".")[-1] if "." in (filename or "") else ""
    if ext == "pdf":
        return "pdf"
    if ext in ("docx", "doc"):
        return "docx"
    if ext == "xlsx":
        return "xlsx"
    if ext == "xls":
        return "txt"  # old Excel; openpyxl only supports xlsx
    if ext == "xer":
        return "xer"  # Primavera P6 export
    if ext in ("jpg", "jpeg", "png", "gif", "webp"):
        return "image"
    if ext in ("txt", "md", "csv", "json", "xml", "html", "htm"):
        return "txt" if ext in ("txt", "md") else ("csv" if ext == "csv" else "txt")
    return "unknown"


def extract_text(bytes_in: bytes, filename: str, format_hint: str | None = None) -> Tuple[str, str]:
    """
    Extract raw text from file bytes. Returns (raw_text, format).
    format_hint can override detection (pdf, docx, xlsx, txt).
    """
    fmt = format_hint or detect_format(filename, bytes_in)
    text = ""
    if fmt == "pdf":
        text = _extract_pdf(bytes_in)
    elif fmt == "docx":
        text = _extract_docx(bytes_in)
    elif fmt == "xlsx":
        text = _extract_xlsx(bytes_in)
    elif fmt == "xer":
        text = _extract_xer(bytes_in)
    elif fmt == "image":
        text = _extract_image(bytes_in, filename)
    elif fmt in ("txt", "csv", "unknown"):
        text = _extract_txt(bytes_in)
    else:
        text = _extract_txt(bytes_in)
    return (text or "").strip(), fmt
